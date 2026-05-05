import io

import PyPDF2
import google.generativeai as genai
from docx import Document
from fastapi import HTTPException

from .config import (
    ALLOWED_FILE_TYPES,
    GEMINI_API_KEY,
    MAX_FILE_SIZE_MB,
    RATE_LIMIT_PER_MINUTE,
    logger,
)


def initialize_model():
    """Initialize Gemini once so the rest of the app can ask for AI content."""
    if not GEMINI_API_KEY:
        logger.warning("GEMINI_API_KEY not found in .env file. AI features will be disabled.")
        logger.warning("Get a key from https://makersuite.google.com/ and add it to your .env file.")
        return None

    try:
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel("gemini-1.5-flash")
        model.generate_content(
            "test",
            generation_config=genai.types.GenerationConfig(max_output_tokens=1),
        )
        logger.info("Gemini AI initialized and tested successfully.")
        return model
    except Exception as exc:
        logger.error("Failed to initialize Gemini AI. Please check your API key and configuration.")
        logger.error(f"Error details: {exc}")
        logger.warning("AI features will be disabled. The application will run in demo mode.")
        return None


class ResumeOptimizer:
    """Owns the domain logic for extracting and generating resume content."""

    def __init__(self, model):
        self.model = model
        self.use_gemini = model is not None

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF uploads."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            return "\n".join(page.extract_text() or "" for page in pdf_reader.pages).strip()
        except Exception as exc:
            logger.error(f"PDF extraction error: {exc}")
            raise HTTPException(status_code=400, detail=f"Error reading PDF: {exc}") from exc

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX uploads."""
        try:
            doc = Document(io.BytesIO(file_content))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()
        except Exception as exc:
            logger.error(f"DOCX extraction error: {exc}")
            raise HTTPException(status_code=400, detail=f"Error reading DOCX: {exc}") from exc

    def generate_ai_content(self, prompt: str) -> str:
        """Generate text through Gemini, or return a useful demo response."""
        try:
            if self.use_gemini:
                response = self.model.generate_content(prompt)
                return response.text

            return f"""
            # AI Service Configuration Required

            This is a demonstration response. To get actual AI-powered optimization:

            1. Obtain a Google Gemini AI API key from: https://makersuite.google.com/
            2. Add it to your .env file as: GEMINI_API_KEY=your_actual_key_here
            3. Restart the server

            Current configuration:
            - Max file size: {MAX_FILE_SIZE_MB}MB
            - Allowed types: {', '.join(ALLOWED_FILE_TYPES)}
            - Rate limit: {RATE_LIMIT_PER_MINUTE} requests/minute

            The system will provide real AI-powered resume optimization once properly configured.
            """
        except Exception as exc:
            logger.error(f"AI generation error: {exc}")
            raise HTTPException(status_code=500, detail=f"AI generation error: {exc}") from exc

    def extract_keywords(self, job_description: str) -> str:
        """Extract ATS keywords that matter for the target role."""
        prompt = f"""
        Analyze the following job description and extract the most important ATS (Applicant Tracking System) keywords and phrases that should be included in a resume. Focus on:
        1. Technical skills and technologies
        2. Required qualifications and certifications
        3. Industry-specific terms
        4. Action verbs and competencies
        5. Job titles and roles mentioned

        Job Description:
        {job_description}

        Please provide a comprehensive list of keywords and phrases that would help a resume pass ATS screening for this position. Format the response as a clean, organized list.
        """
        return self.generate_ai_content(prompt)

    def optimize_resume(self, original_resume: str, job_description: str, keywords: str) -> str:
        """Produce an improved resume while keeping the candidate's facts intact."""
        prompt = f"""
        You are an expert resume writer and ATS optimization specialist. Please optimize the following resume to better match the job description while maintaining truthfulness and the candidate's authentic experience.

        ORIGINAL RESUME:
        {original_resume}

        JOB DESCRIPTION:
        {job_description}

        ATS KEYWORDS TO INTEGRATE:
        {keywords}

        OPTIMIZATION INSTRUCTIONS:
        1. Maintain all factual information about the candidate's experience
        2. Integrate relevant keywords naturally into existing content
        3. Enhance bullet points to better highlight relevant achievements
        4. Improve formatting and structure for ATS compatibility
        5. Quantify achievements where possible
        6. Use strong action verbs
        7. Ensure the resume directly addresses job requirements
        8. Keep the same overall length and structure

        Please return only the optimized resume content, properly formatted and ready to use. Do not include any additional commentary or explanations.
        """
        return self.generate_ai_content(prompt)

    def generate_cover_letter(self, optimized_resume: str, job_description: str) -> str:
        """Generate the matching cover letter for the same application."""
        prompt = f"""
        Based on the following optimized resume and job description, create a compelling, personalized cover letter that:

        1. Addresses the specific role and company
        2. Highlights the most relevant qualifications from the resume
        3. Shows enthusiasm and cultural fit
        4. Follows professional cover letter structure
        5. Is concise but impactful (3-4 paragraphs)
        6. Uses a professional, confident tone

        OPTIMIZED RESUME:
        {optimized_resume}

        JOB DESCRIPTION:
        {job_description}

        Please create a complete cover letter that would accompany this resume. Include placeholders like [Company Name], [Hiring Manager Name], [Your Name] where specific details would need to be customized. Format it as a professional business letter.
        """
        return self.generate_ai_content(prompt)

    def create_docx_from_text(self, content: str, title: str) -> bytes:
        """Turn generated plain text into a DOCX download."""
        doc = Document()
        doc.add_heading(title, 0)
        for paragraph in content.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

