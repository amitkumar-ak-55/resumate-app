from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from datetime import datetime

def create_resume_docx(resume_content: str, output_path: str):
    """
    Create a DOCX file from optimized resume content.
    
    Args:
        resume_content (str): AI-optimized resume text
        output_path (str): Path where the DOCX file should be saved
    """
    try:
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Process the resume content
        lines = resume_content.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Check if this is likely a section header
            if is_section_header(line):
                # Add section header
                if current_section:  # Add space before new section (except first)
                    doc.add_paragraph()
                
                heading = doc.add_paragraph()
                heading_run = heading.add_run(line.upper())
                heading_run.font.size = Pt(12)
                heading_run.font.bold = True
                heading_run.font.color.rgb = None  # Black
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add underline
                heading_run.font.underline = True
                
                current_section = line
                
            elif is_main_header(line, i):
                # This is likely the candidate's name
                name_para = doc.add_paragraph()
                name_run = name_para.add_run(line)
                name_run.font.size = Pt(18)
                name_run.font.bold = True
                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif is_contact_info(line):
                # Contact information
                contact_para = doc.add_paragraph()
                contact_run = contact_para.add_run(line)
                contact_run.font.size = Pt(10)
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            else:
                # Regular content
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.size = Pt(11)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add bullet point if this looks like a list item
                if line.startswith(('•', '-', '*')) or (current_section and 
                    any(keyword in current_section.lower() for keyword in ['experience', 'skills', 'education', 'achievement'])):
                    para.style = 'List Bullet'
        
        doc.save(output_path)
        
    except Exception as e:
        raise Exception(f"Failed to create resume DOCX: {str(e)}")

def create_cover_letter_docx(cover_letter_content: str, output_path: str):
    """
    Create a DOCX file from cover letter content.
    
    Args:
        cover_letter_content (str): AI-generated cover letter text
        output_path (str): Path where the DOCX file should be saved
    """
    try:
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add spacing
        doc.add_paragraph()
        
        # Process cover letter content
        paragraphs = cover_letter_content.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
                
            # Check if this is the greeting
            if para_text.startswith(('Dear', 'To')):
                greeting_para = doc.add_paragraph()
                greeting_run = greeting_para.add_run(para_text)
                greeting_run.font.size = Pt(12)
                greeting_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph()  # Add space after greeting
                
            # Check if this is the closing
            elif para_text.startswith(('Sincerely', 'Best regards', 'Thank you')):
                doc.add_paragraph()  # Add space before closing
                closing_para = doc.add_paragraph()
                closing_run = closing_para.add_run(para_text)
                closing_run.font.size = Pt(12)
                closing_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            else:
                # Regular paragraph
                para = doc.add_paragraph()
                run = para.add_run(para_text)
                run.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = Inches(0.5)
        
        doc.save(output_path)
        
    except Exception as e:
        raise Exception(f"Failed to create cover letter DOCX: {str(e)}")

def is_section_header(line: str) -> bool:
    """Check if a line is likely a section header."""
    section_keywords = [
        'professional summary', 'summary', 'objective', 'profile',
        'experience', 'work experience', 'employment', 'career history',
        'education', 'qualifications', 'academic background',
        'skills', 'technical skills', 'core competencies', 'expertise',
        'achievements', 'accomplishments', 'awards',
        'certifications', 'licenses', 'projects', 'publications',
        'references', 'interests', 'hobbies', 'volunteer'
    ]
    
    line_lower = line.lower().strip()
    return (
        len(line) < 50 and  # Headers are usually short
        any(keyword in line_lower for keyword in section_keywords) or
        (line.isupper() and len(line.split()) <= 3)  # Short uppercase text
    )

def is_main_header(line: str, position: int) -> bool:
    """Check if a line is likely the main header (candidate name)."""
    return (
        position < 3 and  # Usually at the top
        len(line.split()) <= 4 and  # Names are usually short
        not '@' in line and  # Not an email
        not any(char.isdigit() for char in line)  # No numbers
    )

def is_contact_info(line: str) -> bool:
    """Check if a line contains contact information."""
    return (
        '@' in line or  # Email
        re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', line) or  # Phone number
        'linkedin.com' in line.lower() or
        'github.com' in line.lower() or
        any(word in line.lower() for word in ['street', 'ave', 'blvd', 'dr', 'ct', 'ln'])
    )