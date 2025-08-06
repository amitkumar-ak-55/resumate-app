import fitz  # PyMuPDF
from docx import Document
import os

def parse_resume(file_path: str) -> str:
    """
    Parse resume content from PDF or DOCX file.
    
    Args:
        file_path (str): Path to the resume file
        
    Returns:
        str: Extracted text content from the resume
    """
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return parse_pdf(file_path)
        elif file_extension == '.docx':
            return parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        raise Exception(f"Failed to parse resume: {str(e)}")

def parse_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
            text += "\n"  # Add page break
        
        doc.close()
        
        # Clean up the text
        text = clean_text(text)
        
        if not text.strip():
            raise Exception("No readable text found in the PDF file")
            
        return text
        
    except Exception as e:
        raise Exception(f"Failed to parse PDF: {str(e)}")

def parse_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx."""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Clean up the text
        text = clean_text(text)
        
        if not text.strip():
            raise Exception("No readable text found in the DOCX file")
            
        return text
        
    except Exception as e:
        raise Exception(f"Failed to parse DOCX: {str(e)}")

def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        if line:  # Skip empty lines
            lines.append(line)
    
    # Join lines with proper spacing
    cleaned_text = '\n'.join(lines)
    
    # Remove excessive newlines (more than 2)
    while '\n\n\n' in cleaned_text:
        cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
    
    return cleaned_text.strip()